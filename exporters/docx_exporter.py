from io import BytesIO
from datetime import datetime

from docx import Document


def export_docx(data: dict) -> bytes:
    doc = Document()

    # Safe timestamp
    timestamp = data.get("timestamp", "")

    if timestamp:
        try:
            ts = datetime.fromisoformat(
                timestamp.replace("Z", "+00:00")
            ).strftime("%Y-%m-%d %H:%M:%S UTC")
        except ValueError:
            ts = "Unknown"
    else:
        ts = "Unknown"

    doc.add_heading("ToneShift Rewrite", level=1)
    doc.add_paragraph(f"Date: {ts}")
    doc.add_paragraph(f"Audience: {data.get('audience', '')}")
    doc.add_paragraph(f"Tone: {data.get('tone', '')}")

    doc.add_heading("Original Text", level=2)
    for para in data.get("original_text", "").split("\n"):
        if para.strip():
            doc.add_paragraph(para)

    doc.add_heading("Rewritten Text", level=2)
    for para in data.get("rewritten_text", "").split("\n"):
        if para.strip():
            doc.add_paragraph(para)

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()